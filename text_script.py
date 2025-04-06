# import google.generativeai as genai
# import docx
# import json
# import re

# from config import API_KEY, CSE_ID, GSE_API_KEY
# from verification.enhanced_search_system import VerificationAgent
# # Configure the API key
# genai.configure(api_key="AIzaSyA3DyDSN3TucUfxmOfRPJAOT_WGRD_daZs")  # Replace with your actual API key


# # Constants for chunking
# CHUNK_SIZE = 800   # Process 800 words at a time
# OVERLAP = 50       # Overlap 50 words for continuity

# def extract_claims_with_gemini(text):
#     """
#     Calls Gemini API to extract claims, ensuring JSON format.
#     """
#     try:
#         model = genai.GenerativeModel('gemini-2.0-flash')

#         system_prompt = """
#         You are an AI assistant specializing in extracting factual claims from text. Your task is to:
#         - Identify **factual claims** while **avoiding breaking news headlines**.
#         - If a **speech or transcript is detected**, treat it **as a single claim** instead of breaking it up.
#         - **Exclude** opinions, speculations, or direct speech.
#         - Format your output as a **JSON object** with a key `"claims"` that contains a **list of extracted claims**.
#         - Ensure the response **strictly follows JSON syntax** (double quotes `" "`, no trailing commas).
#         - If no claims are found, return `{ "claims": [] }`.
#         """

#         prompt = f"Extract claims from the following text: {text}"

#         response = model.generate_content([system_prompt, prompt])

#         if response and response.text:
#             print(response.text.strip())
#             return response.text.strip()
#         else:
#             return '{ "claims": [] }'

#     except Exception as e:
#         print(f"Error calling Gemini: {e}")
#         return '{ "claims": [] }'

# def read_text_from_file(file_path):
#     """
#     Reads text content from a .txt or .docx file.
#     """
#     try:
#         if file_path.lower().endswith('.txt'):
#             with open(file_path, 'r', encoding='utf-8') as file:
#                 return file.read()
#         elif file_path.lower().endswith('.docx'):
#             doc = docx.Document(file_path)
#             return '\n'.join(paragraph.text for paragraph in doc.paragraphs)
#         else:
#             print("Unsupported file type. Please provide a .txt or .docx file.")
#             return None
#     except FileNotFoundError:
#         print(f"File not found: {file_path}")
#         return None
#     except docx.opc.exceptions.PackageNotFoundError:
#         print(f"Error opening Word document: {file_path}. Is it a valid .docx file?")
#         return None
#     except Exception as e:
#         print(f"Error processing file {file_path}: {e}")
#         return None

# def split_text_into_chunks(text, chunk_size=CHUNK_SIZE, overlap=OVERLAP):
#     """
#     Splits the text into overlapping chunks of `chunk_size` words with `overlap` words repeating.
#     """
#     words = text.split()
#     chunks = []
    
#     start = 0
#     while start < len(words):
#         end = min(start + chunk_size, len(words))
#         chunk = " ".join(words[start:end])
#         chunks.append(chunk)
#         start += chunk_size - overlap  # Move forward while keeping an overlap
    
#     return chunks

# def extract_claims(response_text):
#     """
#     Extracts claims from the JSON response using regex and JSON parsing.
#     """
#     try:
#         # Extract JSON content using regex (handles malformed outputs)
#         json_match = re.search(r'\{.*\}', response_text, re.DOTALL)
        
#         if json_match:
#             json_text = json_match.group()
#             parsed_json = json.loads(json_text)  # Convert string to Python dictionary
            
#             if isinstance(parsed_json, dict) and "claims" in parsed_json:
#                 return parsed_json["claims"]  # Return the list of claims
#             else:
#                 print("Unexpected response format. Returning raw response.")
#                 return [response_text]  # Return as a single-item list if not properly formatted

#     except json.JSONDecodeError:
#         print("Error parsing JSON. Returning raw response.")
#         return [response_text]  # Return as-is if parsing fails
# def extract_claims_from_string(json_string):
#     """
#     Extracts a list of claims from a JSON-formatted string.

#     Args:
#         json_string (str): A JSON string with a top-level 'claims' key.

#     Returns:
#         list: A list of claim strings, or an empty list if parsing fails.
#     """
#     try:
#         data = json.loads(json_string)
#         return data.get("claims", [])
#     except json.JSONDecodeError as e:
#         print("Error decoding JSON:", e)

#         return []
# def extract_verification_result(data, as_list=False):
#     verification = data.get("verification", {})
    
#     result = {
#         "Is Verified": verification.get("is_verified", "Unknown"),
#         "Confidence": verification.get("confidence", "Unknown"),
#         "Reasoning": verification.get("reasoning", "No reasoning provided."),
#         "Relevant Quotes": verification.get("relevant_quotes", "No quotes provided."),
#         "Label": verification.get("label", "No label provided."),
#         "Context": verification.get("context", "No context provided."),
#         "Source Link": verification.get("source_link", "No source link provided.")
#     }
    
#     if as_list:
#         return list(result.values())  # Return as a list
#     else:
#         return "\n".join([f"{key}: {value}" for key, value in result.items()])  # Return as a formatted string
# # Main Execution
# def execute(file_path):
#     """
#     Reads text from a file, processes it in chunks, and extracts claims in structured JSON format.
#     """
#     text_content = read_text_from_file(file_path)
    
#     if text_content:
#         text_chunks = split_text_into_chunks(text_content)
#         all_claims = []

#         for i, chunk in enumerate(text_chunks):
#             print(f"\nProcessing chunk {i + 1}/{len(text_chunks)}...")
#             claims_output = extract_claims_with_gemini(chunk)
            
#             if claims_output and claims_output != '{ "claims": [] }':
#                 extracted_claims = extract_claims(claims_output)  # Convert Gemini output to list
#                 all_claims.extend(extracted_claims)  # Append correctly formatted claims
        
#         # Output as a structured JSON format
#         final_json = json.dumps({"claims": all_claims}, indent=4)
#         print("\nFinal Extracted Claims (JSON Format):")
#         print(final_json)  # Ensures output is structured
        
#         return final_json  # This allows you to use it in further processing

# # Example Usage
# claims_json = execute(r"test.txt")
# match = re.search(r'\{.*\}', claims_json, re.DOTALL)
# if match:
#     json_text = match.group()
#     claims = extract_claims_from_string(json_text)
#     print("Final list ---------------->",claims)
# else:
#     print("No valid JSON found.")

# # print("\nExtracted Claims JSON Ready for Use:", claims_json)
# for i in range(3):
    
#     print("for claim: ",claims[i])
#     agent = VerificationAgent(api_key=API_KEY, cse_id=CSE_ID, gse_api_key=GSE_API_KEY)
#     text_result = agent.process_query(claims[i])
#     text_reasons = extract_verification_result(text_result, as_list=True)
#     print(text_result)
#     print(text_reasons)


import google.generativeai as genai
import docx
import json
import re

from config import API_KEY, CSE_ID, GSE_API_KEY
from verification.enhanced_search_system import VerificationAgent

# Configure the API key for Gemini
genai.configure(api_key="AIzaSyA3DyDSN3TucUfxmOfRPJAOT_WGRD_daZs")  # Replace with your actual API key

# Constants for chunking
CHUNK_SIZE = 800   # Process 800 words at a time
OVERLAP = 50       # Overlap 50 words for continuity

def extract_claims_with_gemini(text):
    """
    Calls Gemini API to extract claims from the provided text.
    Returns the raw text response, which should be in JSON format.
    """
    try:
        model = genai.GenerativeModel('gemini-2.0-flash')

        system_prompt = """
        You are an AI assistant specializing in extracting factual claims from text. Your task is to:
        - Identify **factual claims** while **avoiding breaking news headlines**.
        - If a **speech or transcript is detected**, treat it **as a single claim** instead of breaking it up.
        - **Exclude** opinions, speculations, or direct speech.
        - Format your output as a **JSON object** with a key "claims" that contains a **list of extracted claims**.
        - Ensure the response **strictly follows JSON syntax** (double quotes " ", no trailing commas).
        - If no claims are found, return { "claims": [] }.
        """

        prompt = f"Extract claims from the following text: {text}"
        response = model.generate_content([system_prompt, prompt])

        if response and response.text:
            return response.text.strip()
        else:
            return '{ "claims": [] }'

    except Exception as e:
        print(f"Error calling Gemini: {e}")
        return '{ "claims": [] }'

def read_text_from_file(file_path):
    """
    Reads text content from a .txt or .docx file.
    """
    try:
        if file_path.lower().endswith('.txt'):
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        elif file_path.lower().endswith('.docx'):
            doc = docx.Document(file_path)
            return '\n'.join(paragraph.text for paragraph in doc.paragraphs)
        else:
            print("Unsupported file type. Please provide a .txt or .docx file.")
            return None
    except FileNotFoundError:
        print(f"File not found: {file_path}")
        return None
    except docx.opc.exceptions.PackageNotFoundError:
        print(f"Error opening Word document: {file_path}. Is it a valid .docx file?")
        return None
    except Exception as e:
        print(f"Error processing file {file_path}: {e}")
        return None

def split_text_into_chunks(text, chunk_size=CHUNK_SIZE, overlap=OVERLAP):
    """
    Splits the text into overlapping chunks of `chunk_size` words with an overlap of `overlap` words.
    """
    words = text.split()
    chunks = []
    start = 0
    while start < len(words):
        end = min(start + chunk_size, len(words))
        chunk = " ".join(words[start:end])
        chunks.append(chunk)
        start += chunk_size - overlap  # Move forward while keeping an overlap
    return chunks

def extract_claims(response_text):
    """
    Extracts claims from the JSON response using regex and JSON parsing.
    """
    try:
        # Extract JSON content using regex (handles malformed outputs)
        json_match = re.search(r'\{.*\}', response_text, re.DOTALL)
        if json_match:
            json_text = json_match.group()
            parsed_json = json.loads(json_text)
            if isinstance(parsed_json, dict) and "claims" in parsed_json:
                return parsed_json["claims"]
            else:
                print("Unexpected response format. Returning raw response.")
                return [response_text]
    except json.JSONDecodeError:
        print("Error parsing JSON. Returning raw response.")
    return [response_text]

def extract_claims_from_string(json_string):
    """
    Extracts a list of claims from a JSON-formatted string.
    """
    try:
        data = json.loads(json_string)
        return data.get("claims", [])
    except json.JSONDecodeError as e:
        print("Error decoding JSON:", e)
        return []

def get_verification_result_dict(result_data):
    """
    Converts the verification result from the agent to a dictionary.
    """
    verification = result_data.get("verification", {})
    return {
        "is_verified": verification.get("is_verified", False),
        "confidence": verification.get("confidence", "Unknown"),
        "reasoning": verification.get("reasoning", "No reasoning provided."),
        "relevant_quotes": verification.get("relevant_quotes", "No quotes provided."),
        "label": verification.get("label", "No label provided."),
        "context": verification.get("context", "No context provided."),
        "source_link": verification.get("source_link", "No source link provided.")
    }

def process_and_verify_claims(file_path):
    """
    Processes the provided file by:
      1. Reading the text.
      2. Splitting the text into overlapping chunks.
      3. Extracting claims from each chunk using the Gemini API.
      4. Verifying each claim using the VerificationAgent.
      5. Printing the verification results for the first three claims.
    
    Returns a dictionary with:
      - "claims_verification": A list of dictionaries for each claim containing:
            "claim": The claim text.
            "verification": The full verification result as a dictionary.
      - "manual_verification_required": A dictionary mapping claims to their reasoning if verification is False.
    """
    text_content = read_text_from_file(file_path)
    if not text_content:
        return None

    # Extract claims from the text chunks
    text_chunks = split_text_into_chunks(text_content)
    all_claims = []
    for i, chunk in enumerate(text_chunks):
        print(f"\nProcessing chunk {i + 1}/{len(text_chunks)}...")
        claims_output = extract_claims_with_gemini(chunk)
        if claims_output and claims_output != '{ "claims": [] }':
            extracted_claims = extract_claims(claims_output)
            all_claims.extend(extracted_claims)

    # Remove duplicate claims if any
    all_claims = list(dict.fromkeys(all_claims))

    # Process verification for each claim
    claims_verification = []
    manual_verification_required = {}

    for i in range(3):
        print(f"\nVerifying claim: {all_claims[i]}")
        agent = VerificationAgent(api_key=API_KEY, cse_id=CSE_ID, gse_api_key=GSE_API_KEY)
        result_data = agent.process_query(all_claims[i])
        verification_dict = get_verification_result_dict(result_data)
        
        # Append the claim and its verification result
        
        
        print("Verification Result:", type(verification_dict["is_verified"]))
        # If the claim is not verified, store the reasoning for manual review.
        if verification_dict["is_verified"] == "FALSE":
            manual_verification_required[all_claims[i]] = verification_dict.get("reasoning", "No reasoning provided.")
        else:
            claims_verification.append({
            "claim": all_claims[i],
            "verification": verification_dict
        })

    # Final loop: Print the verification results for the first three claims
    print("\n--- Final Verification Results for the Top 3 Claims ---")
    for i, item in enumerate(claims_verification):
        print(f"\nFor claim {i + 1}: {item['claim']}")
        print("Verification Result:")
        for key, value in item['verification'].items():
            print(f"  {key}: {value}")

    return {
        "claims_verification": claims_verification,
        "manual_verification_required": manual_verification_required
    }

# If running this module directly, you can test the function as follows:
if __name__ == '__main__':
    file_path = r"test.txt"  # Replace with your actual file path
    results = process_and_verify_claims(file_path)
    if results:
        print("\nFinal Claims Verification Results:")
        print(json.dumps(results, indent=4))
    else:
        print("Failed to process the file.")

